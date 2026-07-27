#!/usr/bin/env python3
"""Export Io world design markdown docs to Microsoft 365 formats (.docx, .xlsx)."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter

DESIGN_DIR = Path(__file__).resolve().parent.parent
OUT_DIR = Path(__file__).resolve().parent

DOCX_SOURCES = [
    "Io_World_Content_Executive_Summary.md",
    "Io_World_Content_Phase_Map.md",
    "Io_World_Content_Milestone_Tickets.md",
    "Io_Biome_Ecology_Roster.md",
    "Io_Biome_Exploration_Gameplay_Plan.md",
    "Io_Underground_Architecture_Plan.md",
]

BRAND_DARK_NAVY = RGBColor(0x1C, 0x2A, 0x38)
BRAND_FUCHSIA = RGBColor(0xC0, 0x2E, 0x7A)
BRAND_BODY = RGBColor(0x2F, 0x2F, 0x2F)


def set_document_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = BRAND_BODY
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")

    for level, size in [(1, 20), (2, 16), (3, 13)]:
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = BRAND_DARK_NAVY if level == 1 else BRAND_FUCHSIA


def add_formatted_run(paragraph, text: str) -> None:
    """Parse **bold** and `code` inline markers."""
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos : match.start()])
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def parse_table_row(line: str) -> list[str]:
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [cell.strip() for cell in line.split("|")]


def is_table_separator(line: str) -> bool:
    return bool(re.match(r"^\|?[\s\-:|]+\|?$", line.strip()))


def markdown_to_docx(md_path: Path, docx_path: Path) -> None:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()
    set_document_styles(doc)

    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    i = 0
    in_code = False
    code_buffer: list[str] = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_buffer))
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                code_buffer = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_buffer.append(line)
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        if stripped == "---":
            doc.add_paragraph("—" * 40)
            i += 1
            continue

        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            level = min(level, 3)
            title = stripped[level:].strip()
            doc.add_heading(title, level=level)
            i += 1
            continue

        if stripped.startswith("|") and i + 1 < len(lines) and is_table_separator(lines[i + 1]):
            headers = parse_table_row(stripped)
            i += 2
            rows: list[list[str]] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                if not is_table_separator(lines[i]):
                    rows.append(parse_table_row(lines[i]))
                i += 1
            table = doc.add_table(rows=1 + len(rows), cols=len(headers))
            table.style = "Table Grid"
            for col, header in enumerate(headers):
                cell = table.rows[0].cells[col]
                cell.text = header
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.bold = True
            for r_idx, row in enumerate(rows):
                for c_idx in range(len(headers)):
                    value = row[c_idx] if c_idx < len(row) else ""
                    table.rows[r_idx + 1].cells[c_idx].text = value
            doc.add_paragraph()
            continue

        if stripped.startswith("- [ ]") or stripped.startswith("- [x]"):
            checked = stripped.startswith("- [x]")
            text = stripped[5:].strip()
            p = doc.add_paragraph(style="List Bullet")
            mark = "☑ " if checked else "☐ "
            add_formatted_run(p, mark + text)
            i += 1
            continue

        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_formatted_run(p, stripped[2:])
            i += 1
            continue

        if re.match(r"^\d+\.\s", stripped):
            p = doc.add_paragraph(style="List Number")
            add_formatted_run(p, re.sub(r"^\d+\.\s", "", stripped))
            i += 1
            continue

        p = doc.add_paragraph()
        add_formatted_run(p, stripped)
        i += 1

    doc.save(docx_path)


def extract_milestone_tickets(md_path: Path) -> list[dict[str, str]]:
    text = md_path.read_text(encoding="utf-8")
    chunks = re.split(r"\n### (IO-W\d+-\d+) · ", text)
    tickets: list[dict[str, str]] = []

    # chunks[0] is preamble; then id, body, id, body, ...
    for i in range(1, len(chunks), 2):
        ticket_id = chunks[i]
        body = chunks[i + 1]
        title_line, _, rest = body.partition("\n")
        title = title_line.strip()

        fields = {
            "ID": ticket_id,
            "Title": title,
            "Track": "",
            "Blocked By": "",
            "Unlocks": "",
            "Phase": re.search(r"W\d+", ticket_id).group(0) if re.search(r"W\d+", ticket_id) else "",
            "Description": "",
            "Acceptance Criteria": "",
            "Refs": "",
        }

        for line in rest.splitlines():
            if "**Track**" in line and "|" in line:
                fields["Track"] = line.split("|", 3)[2].strip()
            elif "**Blocked by**" in line and "|" in line:
                fields["Blocked By"] = line.split("|", 3)[2].strip()
            elif "**Unlocks**" in line and "|" in line:
                fields["Unlocks"] = line.split("|", 3)[2].strip()

        desc_match = re.search(r"\*\*Description:\*\* (.+?)(?=\n\n\*\*Acceptance)", rest, re.DOTALL)
        if desc_match:
            fields["Description"] = desc_match.group(1).strip().replace("\n", " ")

        ac_match = re.search(r"\*\*Acceptance criteria:\*\*\n(.*?)(?=\n\n\*\*Refs|\n\n---|\Z)", rest, re.DOTALL)
        if ac_match:
            criteria = [
                ac_line.strip()
                for ac_line in ac_match.group(1).splitlines()
                if ac_line.strip().startswith("- [")
            ]
            fields["Acceptance Criteria"] = "\n".join(criteria)

        refs_match = re.search(r"\*\*Refs:\*\* (.+)", rest)
        if refs_match:
            fields["Refs"] = refs_match.group(1).strip()

        tickets.append(fields)

    return tickets


def tickets_to_xlsx(tickets: list[dict[str, str]], xlsx_path: Path) -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "Milestone Tickets"

    headers = [
        "ID",
        "Phase",
        "Title",
        "Track",
        "Blocked By",
        "Unlocks",
        "Description",
        "Acceptance Criteria",
        "Refs",
        "Status",
    ]
    header_fill = PatternFill("solid", fgColor="1C2A38")
    header_font = Font(color="EDE9E4", bold=True, name="Calibri", size=11)

    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

    for row_idx, ticket in enumerate(tickets, 2):
        for col_idx, key in enumerate(headers, 1):
            value = ticket.get(key, "") if key != "Status" else "Open"
            cell = ws.cell(row=row_idx, column=col_idx, value=value)
            cell.alignment = Alignment(vertical="top", wrap_text=True)
            cell.font = Font(name="Calibri", size=10)

    widths = [14, 8, 42, 22, 24, 24, 48, 52, 28, 10]
    for i, width in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(i)].width = width

    ws.freeze_panes = "A2"
    ws.auto_filter.ref = f"A1:{get_column_letter(len(headers))}{len(tickets) + 1}"
    wb.save(xlsx_path)


def write_readme() -> None:
    readme = OUT_DIR / "README.txt"
    readme.write_text(
        "Dark Matter: Genesis — Io World Design (Microsoft 365 exports)\n"
        "============================================================\n\n"
        "Generated from markdown sources in Assets/_Project/Documentation/Design/.\n"
        "Regenerate: python3 export_to_office365.py\n\n"
        "Files:\n"
        "  *.docx  — Word documents (edit in Microsoft 365 / Word)\n"
        "  Io_World_Content_Milestone_Tickets.xlsx — Excel tracker for IO-W* tickets\n\n"
        "Canonical sources remain the .md files in the parent Design folder.\n",
        encoding="utf-8",
    )


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)

    for name in DOCX_SOURCES:
        md_path = DESIGN_DIR / name
        if not md_path.exists():
            print(f"Skip missing: {md_path}")
            continue
        docx_name = md_path.stem + ".docx"
        docx_path = OUT_DIR / docx_name
        markdown_to_docx(md_path, docx_path)
        print(f"Wrote {docx_path.name}")

    tickets_md = DESIGN_DIR / "Io_World_Content_Milestone_Tickets.md"
    if tickets_md.exists():
        tickets = extract_milestone_tickets(tickets_md)
        xlsx_path = OUT_DIR / "Io_World_Content_Milestone_Tickets.xlsx"
        tickets_to_xlsx(tickets, xlsx_path)
        print(f"Wrote {xlsx_path.name} ({len(tickets)} tickets)")

    write_readme()
    print("Done.")


if __name__ == "__main__":
    main()
